import openpyxl
import docx
from docx.enum.text import WD_COLOR_INDEX

print("Ange filen som du vill hämta data ifrån?")
print("Filen ska vara i samma format som den bifogade filen 'bedömning_progr.xlsx'")
xlInput = input("> ")
wb = openpyxl.load_workbook(xlInput) #laddar in kalkylarket som anges i python-programmet

sheet1 = wb.get_sheet_by_name("Centralt innehåll") #namnger de olika arken till sheet 1,2,3,4
sheet2 = wb.get_sheet_by_name("Kunskapskrav")
sheet3 = wb.get_sheet_by_name("kraven")
sheet4 = wb.get_sheet_by_name("Info")

elev = [sheet2.cell(row = i, column = 1).value for i in range(2,sheet2.max_row + 1)] #en array med alla namn

docI = [[ #en array med objekt där varje objekt innehåller vilken nivå varje person har uppnått i alla de olika kraven
    sheet2.cell(row = i, column = 2).value,
    sheet2.cell(row = i, column = 3).value, 
    sheet2.cell(row = i, column = 4).value,
    sheet2.cell(row = i, column = 5).value,
    sheet2.cell(row = i, column = 6).value,
    sheet2.cell(row = i, column = 7). value,
    sheet2.cell(row = i, column = 8).value,
    sheet2.cell(row = i, column = 9).value,
    sheet2.cell(row = i, column = 10).value,
    sheet2.cell(row = i, column= 11).value]
for i in range(2,sheet2.max_row+1)] #En for-loop som gör att varje array i arrayen får en ny rad att kolla i tabellen (i-värdet ändras).

cenGen = [[ #en array med objekt där varje objekt innehåller vilka krav varje person har genomfört
    sheet1.cell(row = i, column = 2).value,
    sheet1.cell(row = i, column = 3).value,
    sheet1.cell(row = i, column = 4).value,
    sheet1.cell(row = i, column = 5).value,
    sheet1.cell(row = i, column = 6).value,
    sheet1.cell(row = i, column = 7).value,
    sheet1.cell(row = i, column = 8).value,
    sheet1.cell(row = i, column = 9).value,]
for i in range(2,sheet1.max_row)]

cenUpp=[ sheet1.cell(row=1, column=i).value for i in range(2,sheet1.max_column)]

krav = [{ #vi gör en array med objekt, där varje objekt har ett kravs alla olika nivåer.
    "Krav": sheet3.cell(row = i, column = 1).value,
    "E": sheet3.cell(row = i, column = 2).value,
    "C": sheet3.cell(row = i, column = 3).value,
    "A": sheet3.cell(row = i, column = 4).value}
for i in range(2,sheet3.max_row)]

info = { #vi gör ett objekt med informationen som finns i ark 4.
    "lärare": sheet4.cell(row = 1, column = 2).value,
    "Kurs": sheet4.cell(row = 2, column = 2).value,
    "Klass": sheet4.cell(row = 3, column = 2).value}

def createDoc(namn):#funktionen för att skapa dokument, tar in variabeln: elev
    doc = docx.Document() #skapar ett dokument som heter doc
    doc.add_heading("Bedömningsmatris "+info["Kurs"],level = 0)
    doc.add_heading("Klass: "+ info["Klass"], level=2)
    doc.add_heading("Namn: "+ elev[namn], level = 2)
    doc.add_heading("Centralt innehåll", level=1)

    for a in range(0,len(cenUpp)-1): #gör listan med kursens innehåll
        par = doc.add_paragraph()
        par.style = "List Bullet"
        run = par.add_run(cenUpp[a])
        if cenGen[namn][a] ==  "Genomfört": #om den delen som skrivs ut är genomförd så blir den i fet stil.
            run.bold=True

    doc.add_heading("Kunskapskrav", level=1)

    table = doc.add_table(sheet3.max_row, cols=4) #skapar en tabell som är 11 rader lång och 4 columner bred.
    a = table.rows[0]
    a.cells[0].text = "Krav"#de här fyra raderna står för de fasta rubrikerna för tabellen.
    a.cells[1].text = "E"
    a.cells[2].text = "C"
    a.cells[3].text = "A"
    for a in range(1,11): #Här har vi en for-loop som skriver in alla olika krav på alla olika nivåer. 
        b = a-1
        a = table.rows[a] #a-vädet används för att bestämma vilken rad vi redigerar just nu.
        a.cells[0].text = krav[b]["Krav"] #Här skriver vi ut det tredje kunskapskravet på den fjärde raden i tabellen.

        ePar = a.cells[1].paragraphs[0]
        eRun = ePar.add_run(krav[b]["E"]) #Här skriver vi in kunskapskraven för E-nivån.
        if docI[namn][b] == "E": #om just det här kravet är betygsatt "E" så ska texten för kunskapskravsnivån bli grönmarkerad
            font = eRun.font
            font.highlight_color = 4 #färg 4 är grön

        cPar = a.cells[2].paragraphs[0]
        cRun = cPar.add_run(krav[b]["C"]) #Här skriver vi in kunskapskraven för C-nivån.
        if docI[namn][b] == "C": #om just det här kravet är betygsatt "C" så ska texten för kunskapskravsnivån bli grönmarkerad
            font = cRun.font
            font.highlight_color = 4

        aPar = a.cells[3].paragraphs[0]
        aRun = aPar.add_run(krav[b]["A"]) #Här skriver vi in kunskapskraven för A-nivån.
        if docI[namn][b] == "A":#om just det här kravet är betygsatt "A" så ska texten för kunskapskravsnivån bli grönmarkerad
            font = aRun.font
            font.highlight_color = 4

    doc.save(elev[namn]+".docx") #till sist så sparar vi dokumentet

for namn in range(0,len(elev)):
    createDoc(namn)